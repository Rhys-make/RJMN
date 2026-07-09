"""Utilities for reading uploaded resume and JD files."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO


SUPPORTED_EXTENSIONS = {".txt", ".docx", ".pdf"}


def _decode_text_bytes(data: bytes) -> str:
    """Decode text bytes with common encodings used in Chinese documents."""
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_txt(file_obj: BinaryIO) -> str:
    return _decode_text_bytes(file_obj.read())


def _read_docx(file_obj: BinaryIO) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx 依赖，请先执行 pip install python-docx") from exc

    document = Document(file_obj)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_text = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_text.append(" | ".join(cells))

    return "\n".join(paragraphs + table_text)


def _read_pdf(file_obj: BinaryIO) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("缺少 pdfplumber 依赖，请先执行 pip install pdfplumber") from exc

    pages_text = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(text.strip())
    return "\n".join(pages_text)


def read_file_by_suffix(file_obj: BinaryIO, filename: str) -> tuple[str, str | None]:
    """Read txt, docx, or pdf content and return (text, error)."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        return "", f"暂不支持 {suffix or '未知'} 文件类型，请上传 txt、docx 或 pdf 文件。"

    try:
        if suffix == ".txt":
            text = _read_txt(file_obj)
        elif suffix == ".docx":
            text = _read_docx(file_obj)
        else:
            text = _read_pdf(file_obj)

        text = text.strip()
        if not text:
            return "", "文件已读取，但未提取到有效文本内容。"
        return text, None
    except Exception as exc:  # noqa: BLE001 - keep UI stable and show a clear message.
        return "", f"文件读取失败：{exc}"


def read_uploaded_file(uploaded_file) -> tuple[str, str | None]:
    """Read a Streamlit UploadedFile object safely."""
    if uploaded_file is None:
        return "", None
    data = uploaded_file.getvalue()
    return read_file_by_suffix(BytesIO(data), uploaded_file.name)


def read_local_file(file_path: str | Path) -> tuple[str, str | None]:
    """Read a local txt, docx, or pdf file. Useful for tests and samples."""
    path = Path(file_path)
    if not path.exists():
        return "", f"文件不存在：{path}"
    with path.open("rb") as file_obj:
        return read_file_by_suffix(file_obj, path.name)

